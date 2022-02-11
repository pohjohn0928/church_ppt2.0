import os
import glob

from PIL import Image
from docx import Document


def delete_files():
    pptx_files = glob.glob('*.{}'.format('pptx'))
    docx_files = glob.glob('*.{}'.format('docx'))

    for docx_file, pptx_file in zip(docx_files, pptx_files):
        os.remove(docx_file)
        os.remove(pptx_file)


def replace_images():
    im1 = Image.open(r"1.jpeg")
    im1.save('2.png')


if __name__ == '__main__':
    document = Document()

    p = document.add_paragraph()
    run = p.add_run()
    run.text = "123\n5678799"
    font = run.font
    font.bold = True
    font.name = '微軟正黑體'

    p2 = document.add_paragraph()
    run = p2.add_run()
    run.text = "121"
    font = run.font

    p3 = document.add_page_break()
    run = p3.add_run()
    run.text = "11111111111111111"
    font = run.font

    document.save(f'123.docx')
