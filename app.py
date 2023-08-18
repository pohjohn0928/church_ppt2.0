import glob

from PIL import Image


from docx import Document
from flask import Flask, request, render_template, redirect, url_for, send_file, make_response, Blueprint
from Helpers.datahelper import ReadPdfFile, MakePPT
from Helpers.docx import Word
from Helpers.email import Gmail
from Helpers.worship_song import make_send_worship_songs
import time
from Helpers.Books import bible_config
import datetime
from time import gmtime, strftime
import os
import json

from test2 import get_file_by_pattern

calvary_ppt = Blueprint('calvary_ppt', __name__)


@calvary_ppt.route('/')
def home():
    resp = make_response(render_template('init.html'))
    resp.delete_cookie('user')
    resp.delete_cookie('account')
    resp.delete_cookie('password')
    return resp


@calvary_ppt.route('/init', methods=["POST", "GET"])
def init():
    ip_address = request.remote_addr
    now = strftime("%Y-%m-%d %H:%M:%S", gmtime())
    print(f'{now}')
    print(f'ip: {ip_address}')
    if request.method == "POST":
        account = request.values['account']
        password = request.values['password']
        if account == 'church_ppt' and password == 'churchchurch':
            resp = make_response(render_template('index.html'))
            resp.set_cookie('user', ip_address)
            resp.set_cookie('account', 'church_ppt')
            resp.set_cookie('password', 'churchchurch')
            return resp
        else:
            return render_template("init.html", message='wrong account or password')
    else:
        account = request.cookies.get('account')
        password = request.cookies.get('password')
        if account == 'church_ppt' and password == 'churchchurch':
            resp = make_response(render_template('index.html'))
            resp.set_cookie('user', ip_address)
            resp.set_cookie('account', 'church_ppt')
            resp.set_cookie('password', 'churchchurch')
            return resp


@calvary_ppt.route('/bible_info', methods=["POST"])
def bible_info():
    bible = bible_config.passage_data
    return bible


@calvary_ppt.route('/getPdfFile', methods=["POST"])
def getPdfFile():
    sermonTitle = request.values['sermon_title']
    closingSongName = request.values['closing_song']
    sr_info = request.form.get('sr_info')
    sr_version = request.form.get('sr_version')
    sis_info = request.form.get('sis_info')
    sis_version = request.form.get('sis_version')
    annocement = request.form.get('annocement')
    receivers = request.form.get('receivers')

    sr_info = json.loads(sr_info)["sr_info"]
    sr_version = json.loads(sr_version)["sr_version"]
    sis_info = json.loads(sis_info)["sis_info"]
    sis_version = json.loads(sis_version)["sis_version"]
    annocement = json.loads(annocement)["annocement"]
    receivers = json.loads(receivers)["receivers"]

    readPdfFile = ReadPdfFile()
    chineseScrpitureReading = readPdfFile.getChieseScripture(sr_info)
    chineseScrpitureInSermon = readPdfFile.getChieseScripture(sis_info)
    englishScrpitureReading = {"verses": sr_info, "bibleVersion": sr_version}
    englishScrpitureInSermon = {"verses": sis_info, "bibleVersion": sis_version}

    closingSong = readPdfFile.getClosingSong(closingSongName)
    blessing_song = readPdfFile.getBlessingSong()

    d = datetime.date.today()
    while d.weekday() != 6:
        d += datetime.timedelta(1)
    date = d

    data = {"annocement": annocement, "englishScrpitureReading": englishScrpitureReading,
            "chineseScrpitureReading": chineseScrpitureReading, "englishScrpitureInSermon": englishScrpitureInSermon,
            "chineseScrpitureInSermon": chineseScrpitureInSermon, "sermonTitle": sermonTitle, "date": date,
            "closingSongName": closingSongName, "closingSong": closingSong, "blessing_song": blessing_song}

    print('englishScrpitureReading : ', englishScrpitureReading["verses"])
    print('chineseScrpitureReading : ', chineseScrpitureReading)
    print('englishScrpitureInSermon : ', englishScrpitureInSermon["verses"])
    print('chineseScrpitureInSermon : ', chineseScrpitureInSermon)

    if len(englishScrpitureReading["verses"]) != len(chineseScrpitureReading):
        return 'Error'
    elif len(englishScrpitureInSermon["verses"]) != len(chineseScrpitureInSermon):
        return 'Error'

    else:
        start = time.time()
        threads = []
        threads.append(MakePPT(data))  # make ppt
        threads[0].start()
        threads.append(Word(data))  # make word
        threads[-1].start()

        for t in threads:
            t.join()
        end = time.time()
        print(f'total cost for ppt and docx : {end - start} sec')

        # Send Email
        start = time.time()
        for receiver in receivers:
            gmail = Gmail()
            base = os.path.dirname(os.path.abspath(__file__))
            base = base.replace('\\', '/')
            gmail.send(receiver, f'Scripture for {data["date"]}', f'{base}/ppt/churchPPT{data["date"]}.pptx', data["sermonTitle"], data["date"])
            end = time.time()
        print("Send Mail cost：%f sec" % (end - start))

        # Done
        return f"PPT Path : {os.path.dirname(__file__)}"


@calvary_ppt.route('/scriptures')
def scriptures():
    return render_template('scriptures.html')


@calvary_ppt.route('/scriptures_file')
def scriptures_file():
    d = datetime.date.today()
    while d.weekday() != 6:
        d += datetime.timedelta(1)
    date = d

    root_path = "/".join(os.getcwd().split('/'))
    document = Document(f"{root_path}/docx/Scripture_In_Sermon{date}.docx")
    return_dict = {'scriptures': [], 'verses': []}
    for i, p in enumerate(document.paragraphs):
        if i == 0:
            return_dict['date'] = p.text
            continue
        else:
            if i % 2 == 1:
                return_dict['scriptures'].append(p.text)
            else:
                return_dict['verses'].append(p.text)
    return return_dict


@calvary_ppt.route('/service-worker.js')
def sw():
    return app.send_static_file('service-worker.js')


@calvary_ppt.route('/edit_announcement.html')
def edit_announcement():
    account = request.cookies.get('account')
    password = request.cookies.get('password')
    if account == 'church_ppt' and password == 'churchchurch':
        return render_template('edit_announcement.html')


@calvary_ppt.route('/announcement_info')
def get_announcement_info():
    account = request.cookies.get('account')
    password = request.cookies.get('password')
    if account == 'church_ppt' and password == 'churchchurch':
        base = os.path.dirname(os.path.abspath(__file__))
        result = get_file_by_pattern(f'{base}/static/annocement', 'png')
        return {"announcements": result}


@calvary_ppt.route('/announcement_info', methods=["PUT"])
def update_announcement_info():
    file = request.files.get('file')
    name = request.values.get('name')
    img = Image.open(file)
    img.save('static/annocement/' + name)
    return "Done!"


@calvary_ppt.route('/worship_songs.html')
def worship_songs():
    account = request.cookies.get('account')
    password = request.cookies.get('password')
    if account == 'church_ppt' and password == 'churchchurch':
        return render_template('worship_songs.html')


@calvary_ppt.route('/worship_song_set')
def worship_song_set():
    account = request.cookies.get('account')
    password = request.cookies.get('password')
    if account == 'church_ppt' and password == 'churchchurch':
        base = os.path.dirname(os.path.abspath(__file__))
        result = get_file_by_pattern(f'{base}/Helpers/worship_songs', 'docx')
        return {"worship_songs": result}


@calvary_ppt.route('/worship_song_lyrics/<song>')
def worship_song_lyrics(song):
    account = request.cookies.get('account')
    password = request.cookies.get('password')
    if account == 'church_ppt' and password == 'churchchurch':
        document = Document(f'Helpers/worship_songs/{song}.docx')
        lyrics = []
        for i, p in enumerate(document.paragraphs):
            lyrics.append(p.text)
        return {"lyrics": lyrics}


@calvary_ppt.route('/make_worship_song_file', methods=["POST"])
def make_worship_song_file():
    account = request.cookies.get('account')
    password = request.cookies.get('password')
    if account == 'church_ppt' and password == 'churchchurch':
        worship_songs = request.form.get('worship_songs')
        worship_songs = json.loads(worship_songs)["worship_songs"]
        receivers = request.form.get('receivers')
        receivers = json.loads(receivers)["receivers"]
        make_send_worship_songs(worship_songs, receivers)
        return "done"


app = Flask(__name__)
app.register_blueprint(calvary_ppt, url_prefix='/')
app.config['JSON_SORT_KEYS'] = False


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=80, debug=False)
    # pyinstaller -w -F --add-data "templates:templates" --add-data "static:static" --add-data "Helpers:Helpers" app.py
