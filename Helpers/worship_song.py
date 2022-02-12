import datetime
from docx.shared import Pt

from Helpers.email import Gmail
from docx import Document


def make_send_worship_songs(songs, receivers):

    d = datetime.date.today()
    while d.weekday() != 6:
        d += datetime.timedelta(1)
    date = d

    worship_sheet = Document()
    para = worship_sheet.add_paragraph()
    for song in songs:
        document = Document(f'Helpers/worship_songs/{song}.docx')
        for i, p in enumerate(document.paragraphs):
            run = para.add_run()
            text = p.text
            run.text = text + '\n'
            font = run.font
            if i == 0:
                font.bold = True
                font.size = Pt(16)
            elif text.lower().startswith('verse') or text.lower().startswith('chorus') or text.lower().startswith('bridge') or text.lower().startswith('pre-chorus') or text.lower().startswith('tag'):
                font.bold = True
            else:
                font.size = Pt(12)
            font.name = 'Courier New'
            # print(p.text)
        # print('----------------------------------------------')
        para = worship_sheet.add_page_break()
    worship_sheet.save(f'worship_songs/{date}.docx')

    for receiver in receivers:
        gmail = Gmail()
        gmail.send_worship_songs(receiver, f'Worship Songs for {date}', f'worship_songs/{date}.docx')

