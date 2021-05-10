from docx import Document
from docx.shared import Inches,Pt,RGBColor
from Helpers.datahelper import MakePPT
import time
import threading

class Word(threading.Thread):   #threading.Thread
    def __init__(self,data):
        threading.Thread.__init__(self)
        self.start_time = time.time()
        self.data = data
        self.document = Document()
        self.date = data["date"]
        self.document.add_heading(f'Scripture for {self.date}', 0)

    def run(self):
        for index in range(len(self.data["englishScrpitureInSermon"]["verses"])):
            scrpiture = self.data["englishScrpitureInSermon"]["verses"][index]
            bible_version = self.data["englishScrpitureInSermon"]["bibleVersion"][index]
            verse = MakePPT(self.data).getEnglishBibleVerses(scrpiture, bible_version)
            scrpiture = self.scrpiture_detail(scrpiture)
            if bible_version != 'NKJV' and bible_version != 'nkjv':
                scrpiture += f" ({bible_version})"
            self.setTitle(scrpiture)
            self.setVerse(verse)
            scrpiture = self.data["chineseScrpitureInSermon"][index]
            verse = MakePPT(self.data).getChineseBibleVerses(scrpiture)
            scrpiture = self.scrpiture_detail(scrpiture)
            self.setTitle(scrpiture, 'chinese')
            self.setVerse(verse, 'chinese')
        self.document.save('Scripture_In_Sermon.docx')
        end_time = time.time()
        print("Make docx cost：%f sec" % (end_time - self.start_time))

    # def addPage(self):
    #     self.document.add_page_break()

    def setTitle(self,Scipture,language = 'english'):
        p = self.document.add_paragraph()
        run = p.add_run()
        run.text = Scipture
        font = run.font
        font.bold = True
        if language != 'english':
            font.name = '微軟正黑體'
        font.color.rgb = RGBColor(255, 0, 0)

    def setVerse(self,verse,language = 'english'):
        p = self.document.add_paragraph()
        for v in verse:
            v = v.replace('<sv>','')
            run = p.add_run()
            run.text = v
            font = run.font
            if language != 'english':
                font.name = '微軟正黑體'

    def scrpiture_detail(self,scrpiture):
        start_verse = scrpiture.split(':')[1].split('-')[0]
        end_verse = scrpiture.split(':')[1].split('-')[1].split(' ')[0]
        if start_verse == end_verse:
            book = scrpiture.split(' ')[0]
            chapter = scrpiture.split(' ')[1].split(':')[0]
            return f'{book} {chapter}:{start_verse}'
        else:
            return scrpiture